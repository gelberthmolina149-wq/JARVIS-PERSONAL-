import os
import json
from pathlib import Path
from datetime import datetime

def document_creator(parameters: dict, player=None) -> str:
    """
    Crea documentos de texto, Word o Excel en base a los parámetros.
    """
    action = parameters.get("action", "").lower()
    title = parameters.get("title", "Documento_Sin_Titulo")
    content = parameters.get("content", "")
    sheets = parameters.get("sheets", [])
    save_path_str = parameters.get("save_path", "").strip()
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join([c for c in title if c.isalpha() or c.isdigit() or c==' ']).rstrip().replace(" ", "_")
    if not safe_title:
        safe_title = "Documento"
        
    # Deserializar sheets si viene como string JSON
    if (action in ("excel", "google_sheet")) and isinstance(sheets, str):
        try:
            sheets = json.loads(sheets)
        except Exception:
            pass

    # Determinar ruta de guardado
    file_path = None
    if save_path_str:
        try:
            path_obj = Path(save_path_str)
            # Si es un directorio (o no tiene extensión), creamos el directorio y agregamos nombre formateado
            if path_obj.is_dir() or not path_obj.suffix:
                path_obj.mkdir(parents=True, exist_ok=True)
                ext = ".docx" if action in ("word", "google_doc") else (".xlsx" if action in ("excel", "google_sheet") else ".txt")
                file_path = path_obj / f"{safe_title}_{timestamp}{ext}"
            else:
                path_obj.parent.mkdir(parents=True, exist_ok=True)
                file_path = path_obj
        except Exception:
            pass

    if not file_path:
        desktop_path = Path(os.path.join(os.environ["USERPROFILE"], "Desktop"))
        ext = ".docx" if action in ("word", "google_doc") else (".xlsx" if action in ("excel", "google_sheet") else ".txt")
        file_path = desktop_path / f"{safe_title}_{timestamp}{ext}"
        
    try:
        if action == "word" or action == "google_doc":
            # Si piden google doc, por ahora lo hacemos Word local y avisamos.
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(title, 0)
                
                # Procesamiento simple de markdown a Word
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        doc.add_paragraph(line)
                        
                file_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(file_path)
                return f"Documento Word creado exitosamente en: '{file_path}'."
            except ImportError:
                return "Error: Faltan librerías para crear Word. (python-docx)"
                
        elif action == "excel" or action == "google_sheet":
            try:
                from openpyxl import Workbook
                wb = Workbook()
                wb.remove(wb.active) # Remove default sheet
                
                if not sheets:
                    return "Error: No se proporcionaron datos (sheets) para crear el Excel."
                    
                for sheet_data in sheets:
                    sheet_name = sheet_data.get("name", "Hoja")
                    headers = sheet_data.get("headers", [])
                    rows = sheet_data.get("rows", [])
                    
                    ws = wb.create_sheet(title=sheet_name[:31])
                    if headers:
                        ws.append(headers)
                    for row in rows:
                        ws.append(row)
                        
                file_path.parent.mkdir(parents=True, exist_ok=True)
                wb.save(file_path)
                return f"Planilla Excel creada exitosamente en: '{file_path}'."
            except ImportError:
                return "Error: Faltan librerías para crear Excel. (openpyxl)"
                
        elif action == "text":
            file_path.parent.mkdir(parents=True, exist_ok=True)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(f"{title}\n\n{content}")
            return f"Archivo de texto creado exitosamente en: '{file_path}'."
            
        else:
            return f"Acción '{action}' no soportada o desconocida. Usá 'word', 'excel' o 'text'."
            
    except Exception as e:
        return f"Error al crear el documento: {str(e)}"
